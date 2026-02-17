# import fitz
# from docx import Document

# def load_document(uploaded_file):
#     if uploaded_file.name.endswith(".pdf"):
#         doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#         text = ""
#         for page in doc:
#             text += page.get_text()
#         return text

#     elif uploaded_file.name.endswith(".docx"):
#         doc = Document(uploaded_file)
#         return "\n".join([p.text for p in doc.paragraphs])

#     else:
#         return uploaded_file.read().decode("utf-8")


import pdfplumber
from docx import Document
import io


def load_document(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()

    if file_type == "pdf":
        return _load_pdf(uploaded_file)

    elif file_type == "docx":
        return _load_docx(uploaded_file)

    elif file_type == "txt":
        return _load_txt(uploaded_file)

    else:
        raise ValueError("Unsupported file format")


def _load_pdf(uploaded_file):
    pages = []

    with pdfplumber.open(uploaded_file) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append({
                    "page": i + 1,
                    "text": text
                })

    return pages


def _load_docx(uploaded_file):
    doc = Document(io.BytesIO(uploaded_file.read()))
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    return [{
        "page": 1,
        "text": text
    }]


def _load_txt(uploaded_file):
    text = uploaded_file.read().decode("utf-8")

    return [{
        "page": 1,
        "text": text
    }]
